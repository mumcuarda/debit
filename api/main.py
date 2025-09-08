from __future__ import annotations
from pathlib import Path
from typing import Dict, Any, Optional, Tuple
import re, io, zipfile
from datetime import datetime
from dateutil.parser import parse as parse_date
from dateutil.relativedelta import relativedelta
from docx import Document
from docxtpl import DocxTemplate

# ----------------------------
# Helpers
# ----------------------------

LABELS = [
    "UNIQUE MARKET REFERENCE",
    "TYPE",
    "INSURED",
    "ADDITINONAL INSURED",
    "ADDRESS (of Reinsured)",
    "REINSURED",
    "PERIOD",
    "PAYMENT TERMS",
    "PREMIUM",
    "TOTAL BROKERAGE",
    "SIGNED LINE",
]

def _doc_text(doc: Document) -> str:
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t: parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            rowtxt = " ".join(c.text.strip() for c in row.cells if c.text.strip())
            if rowtxt:
                parts.append(rowtxt)
    return "\n".join(parts)

def _find(pattern: str, text: str, flags=re.IGNORECASE|re.DOTALL, required=True, group=1) -> Optional[str]:
    m = re.search(pattern, text, flags)
    if not m:
        if required:
            raise ValueError(f"Missing required field using pattern: {pattern}")
        return None
    val = m.group(group).strip()
    val = re.sub(r"\s+", " ", val)
    return val

def _block_after(label: str, text: str, next_labels=LABELS) -> Optional[str]:
    # Capture content after LABEL: up to next ALLCAPS label we know about
    pat = rf"{re.escape(label)}\s*:?\s*(.+?)(?=\n(?:{'|'.join(map(re.escape, next_labels))})\b|$)"
    return _find(pat, text, required=False, group=1)

def _leftmost_date(s: str) -> datetime:
    m = re.search(r"(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})", s)
    if m:
        return parse_date(m.group(1), dayfirst=True)
    return parse_date(s, dayfirst=True, fuzzy=True)

def _parse_currency_amount(raw: str) -> tuple[str, float]:
    # Detect currency from the token next to PREMIUM value (guide)
    cur_m = re.search(r"\b([A-Z]{3}|€|\$|£|TL|TRY)\b", raw)
    cur = {"€":"EUR","$":"USD","£":"GBP","TL":"TRY"}.get(cur_m.group(1), cur_m.group(1)) if cur_m else "EUR"
    amt_m = re.search(r"([\d.,]+)", raw)
    amt = amt_m.group(1) if amt_m else "0"
    # Normalize EU style "50.000,00" -> "50000.00"
    amt_norm = amt.replace(".", "").replace(",", ".")
    try:
        val = float(amt_norm)
    except ValueError:
        val = 0.0
    return cur, val

def _fmt_amount(value: float) -> str:
    # 12.345,67 formatting (as your samples show)
    s = f"{value:,.2f}"
    return s.replace(",", "_").replace(".", ",").replace("_", ".")

def _iban_for_currency(cur: str) -> str:
    # Guide mapping
    return {
        "USD": "TR92 0006 2000 3560 0009 0742 54",
        "EUR": "TR22 0006 2000 3560 0009 0742 53",
    }.get(cur, "")

# ----------------------------
# Parsing — per your helper guide
# ----------------------------

def parse_slip(slip_path: str | Path) -> Dict[str, Any]:
    """
    Extracts:
      - type, insured, reinsured, period, payment terms days, currency & premium
      - UMR (Unique Market Reference), reinsurer name, brokerage%
      - ADDRESS (of Reinsured) -> split into two lines (rule: end right before letter 'N')
    """
    doc = Document(str(slip_path))
    text = _doc_text(doc)

    slip_no = _find(r"\bUNIQUE\s+MARKET\s+REFERENCE\s*:?\s*([^\n]+)", text)  # {{ slip_no }} :contentReference[oaicite:3]{index=3}
    type_   = _find(r"\bTYPE\s*:?\s*([^\n]+)", text)                          # {{ type }}   :contentReference[oaicite:4]{index=4}
    insured = _find(r"\bINSURED\s*:?\s*([^\n]+)", text)                       # {{ insured }}:contentReference[oaicite:5]{index=5}
    reinsured = _find(r"\bREINSURED\s*:?\s*([^\n]+)", text)                   # {{ reinsured }} :contentReference[oaicite:6]{index=6}
    period = _find(r"\bPERIOD\s*:?\s*([^\n]+)", text)                         # {{ period }}  :contentReference[oaicite:7]{index=7}
    premium_raw = _find(r"\bPREMIUM\s*:?\s*([^\n]+)", text)                   # amount/premium base :contentReference[oaicite:8]{index=8}
    currency, premium_val = _parse_currency_amount(premium_raw)               # {{ currency }} from PREMIUM :contentReference[oaicite:9]{index=9}

    # PAYMENT TERMS → days
    terms_line = _find(r"\bPAYMENT\s+TERMS\s*:?\s*([^\n]+)", text)            # number of days :contentReference[oaicite:10]{index=10}
    days_m = re.search(r"(\d{1,3})\s*days", terms_line, flags=re.IGNORECASE)
    days = int(days_m.group(1)) if days_m else 120

    # Due date: leftmost date in PERIOD + days (guide)
    inception = _leftmost_date(period)
    due_date = inception + relativedelta(days=days)                            # {{ term }} :contentReference[oaicite:11]{index=11}

    # ADDRESS (of Reinsured) → split around first 'N' (guide)
    addr_block = _block_after("ADDRESS (of Reinsured)", text, LABELS)          # {{ recipient.*_address_* }} :contentReference[oaicite:12]{index=12}
    addr_block = (addr_block or "").strip()
    first_line, second_line = "", ""
    if addr_block:
        # End right before the letter 'N', rest after 'N'
        idx = addr_block.find("N")
        if idx > 0:
            first_line = addr_block[:idx].rstrip()
            second_line = addr_block[idx:].lstrip()
        else:
            # fallback: split on newline or mid-length
            parts = addr_block.splitlines()
            if len(parts) >= 2:
                first_line, second_line = parts[0], " ".join(parts[1:])
            else:
                mid = len(addr_block)//2
                first_line, second_line = addr_block[:mid], addr_block[mid:]

    # Reinsurer name from SIGNED LINE (strict words before numbers)
    reinsurer = _find(r"\bSIGNED LINE\b.*?\n\s*([A-Z].*?)(?=\s+\d+%|\s+\d{1,3}\s*%|\s*$)", text, required=False) or ""  # {{ reinsurer }} :contentReference[oaicite:13]{index=13}

    # Brokerage percent from TOTAL BROKERAGE → e.g., "20%"
    brokerage_line = _find(r"\bTOTAL\s+BROKERAGE\s*:?\s*([^\n]+)", text, required=False)
    brokerage_pct = 0.0
    if brokerage_line:
        m = re.search(r"(\d{1,2}(?:\.\d+)?)\s*%", brokerage_line)
        if m:
            brokerage_pct = float(m.group(1)) / 100.0

    return {
        "slip_no": slip_no,
        "type": type_,
        "insured": insured,
        "reinsured": reinsured,
        "period": period,
        "currency": currency,
        "premium_value": premium_val,     # numeric
        "premium_raw": premium_raw,       # original string
        "payment_terms_days": days,
        "due_date": due_date.strftime("%d.%m.%Y"),
        "reinsurer": reinsurer,
        "addr_line1": first_line,
        "addr_line2": second_line,
        "brokerage_pct": brokerage_pct,
    }

# ----------------------------
# Context building for Template A & B
#   - Matches your template placeholders exactly.
# ----------------------------

def build_contexts(parsed: Dict[str, Any], reference_a_suffix: str, reference_b_suffix: str) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    today = datetime.now().strftime("%d.%m.%Y")

    # Template A mapping (uses {{ reference_a }}, {{ amount }}, IBAN by currency)  :contentReference[oaicite:14]{index=14}
    ctx_a = {
        # Shared/general
        "date": today,
        "type": parsed["type"],
        "slip_no": parsed["slip_no"],
        "reinsured": parsed["reinsured"],
        "insured": parsed["insured"],
        "period": parsed["period"],
        "term": parsed["due_date"],
        "currency": parsed["currency"],
        "iban": _iban_for_currency(parsed["currency"]),
        # Amount for A is the raw PREMIUM value per guide (“{{ amount }} - PREMIUM”)  :contentReference[oaicite:15]{index=15}
        "amount": _fmt_amount(parsed["premium_value"]),
        # Recipient A (reinsured) + address split
        "recipient.a": parsed["reinsured"],  # guide uses dot notation  :contentReference[oaicite:16]{index=16}
        "recipient.a_address_line_1": parsed["addr_line1"],
        "recipient.a_address_line_2": parsed["addr_line2"],
        # Also provide the flat keys because template_a.docx uses them ({{ recipient_a_address_line_1 }})  :contentReference[oaicite:17]{index=17}
        "recipient_a_address_line_1": parsed["addr_line1"],
        # Reference A (with fixed DN-RHB prefix per guide)
        "reference_a": f"DN-RHB-{reference_a_suffix}",  # :contentReference[oaicite:18]{index=18}
    }

    # Template B mapping (uses {{ reference_b }}, {{ premium }} = premium minus brokerage, and {{ reinsurer }})  :contentReference[oaicite:19]{index=19}
    net_premium = parsed["premium_value"] * (1.0 - parsed["brokerage_pct"])  # 100% - TOTAL BROKERAGE :contentReference[oaicite:20]{index=20}
    ctx_b = {
        "date": today,
        "type": parsed["type"],
        "slip_no": parsed["slip_no"],
        "reinsured": parsed["reinsured"],
        "insured": parsed["insured"],
        "period": parsed["period"],
        "term": parsed["due_date"],
        "premium": _fmt_amount(net_premium),         # {{ premium }} (Due to {{ reinsurer }}) :contentReference[oaicite:21]{index=21}
        "reinsurer": (parsed["reinsurer"] or "").strip(),  # words before numbers from SIGNED LINE :contentReference[oaicite:22]{index=22}
        # Recipient B placeholders — keep placeholders per guide for now (xxx/yyy/zzz)  :contentReference[oaicite:23]{index=23}
        "recipient.b": "xxx",
        "recipient.b_address_line_1": "yyy",
        "recipient.b_address_line_2": "zzz",
        # Reference B
        "reference_b": f"DN-RHB-{reference_b_suffix}",
    }

    return ctx_a, ctx_b

# ----------------------------
# Rendering (two templates)
# ----------------------------

def render_docs(template_a: str|Path, template_b: str|Path, ctx_a: Dict[str, Any], ctx_b: Dict[str, Any]) -> Tuple[bytes, bytes]:
    # Template A
    t = DocxTemplate(str(template_a))
    t.render(ctx_a)
    buf_a = io.BytesIO()
    t.save(buf_a)

    # Template B
    t = DocxTemplate(str(template_b))
    t.render(ctx_b)
    buf_b = io.BytesIO()
    t.save(buf_b)

    return buf_a.getvalue(), buf_b.getvalue()

def zip_outputs(a_bytes: bytes, b_bytes: bytes, name_a: str, name_b: str) -> bytes:
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"{name_a}.docx", a_bytes)
        z.writestr(f"{name_b}.docx", b_bytes)
    return out.getvalue()
